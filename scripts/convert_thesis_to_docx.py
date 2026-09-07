import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

FONT_NAME = "TH SarabunPSK"
FONT_SIZE_BODY = Pt(16)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(16)
FONT_SIZE_TITLE = Pt(20)
FONT_SIZE_CODE = Pt(13)
FONT_SIZE_TABLE = Pt(14)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_run(run, font_name=FONT_NAME, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure East Asia / Complex Script font is set
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}" w:eastAsia="{font_name}"/>')
    rPr.append(rFonts)

def add_formatted_text(paragraph, text, default_size=FONT_SIZE_BODY, default_bold=False, default_italic=False, default_color=None):
    # Regex for bold-italic (***text***), bold (**text**), italic (*text*), inline code (`text`)
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            format_run(run, size=default_size, bold=True, italic=True, color=default_color)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            format_run(run, size=default_size, bold=True, italic=default_italic, color=default_color)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            format_run(run, size=default_size, bold=default_bold, italic=True, color=default_color)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            format_run(run, font_name="Consolas", size=Pt(13), bold=default_bold, color=RGBColor(30, 41, 59))
        else:
            run = paragraph.add_run(part)
            format_run(run, size=default_size, bold=default_bold, italic=default_italic, color=default_color)

def setup_page_setup(section):
    section.top_margin = Inches(1.5)
    section.left_margin = Inches(1.5)
    section.bottom_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)

def convert_markdown_to_docx(md_content, doc=None):
    if doc is None:
        doc = docx.Document()
        setup_page_setup(doc.sections[0])

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code / Diagram blocks (```)
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "F8F9FA")
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                
                # Border box
                tcPr = cell._tc.get_or_add_tcPr()
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>\n'
                    f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>\n'
                    f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="0284C7"/>\n'
                    f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>\n'
                    f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>\n'
                    f'</w:tcBorders>'
                )
                tcPr.append(borders)
                
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.line_spacing = 1.05
                full_code_text = '\n'.join(code_lines)
                run = cp.add_run(full_code_text)
                format_run(run, font_name="Consolas", size=FONT_SIZE_CODE, color=RGBColor(30, 41, 59))
                
                # Add small space after table
                post_p = doc.add_paragraph()
                post_p.paragraph_format.space_after = Pt(6)
                
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle Markdown Tables
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table rows
            parsed_rows = []
            for tl in table_lines:
                # remove leading/trailing pipe and split
                row_cells = [c.strip() for c in tl.strip('|').split('|')]
                # skip divider row like |---|---|
                if all(re.match(r'^:?-+:?$', c) for c in row_cells if c):
                    continue
                parsed_rows.append(row_cells)
            
            if parsed_rows:
                num_cols = max(len(r) for r in parsed_rows)
                num_rows = len(parsed_rows)
                tbl = doc.add_table(rows=num_rows, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl)
                
                for r_idx, r_data in enumerate(parsed_rows):
                    is_header = (r_idx == 0)
                    for c_idx in range(num_cols):
                        val = r_data[c_idx] if c_idx < len(r_data) else ""
                        c = tbl.cell(r_idx, c_idx)
                        set_cell_margins(c, top=80, bottom=80, left=120, right=120)
                        if is_header:
                            set_cell_background(c, "F1F5F9")
                        p = c.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.15
                        add_formatted_text(p, val, default_size=FONT_SIZE_TABLE, default_bold=is_header)
                
                post_tbl_p = doc.add_paragraph()
                post_tbl_p.paragraph_format.space_after = Pt(6)
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, stripped[2:].strip(), default_size=FONT_SIZE_TITLE, default_bold=True)
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, stripped[3:].strip(), default_size=FONT_SIZE_H1, default_bold=True)
            i += 1
            continue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, stripped[4:].strip(), default_size=FONT_SIZE_H2, default_bold=True)
            i += 1
            continue
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, stripped[5:].strip(), default_size=FONT_SIZE_H3, default_bold=True)
            i += 1
            continue

        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, stripped[2:].strip())
            i += 1
            continue

        # Numbered lists (e.g. 1. , 2. )
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(f"{num_match.group(1)}. ")
            format_run(run, bold=True)
            add_formatted_text(p, num_match.group(2))
            i += 1
            continue

        # Normal body paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped)
        i += 1

    return doc

print("Markdown to DOCX converter module loaded.")
