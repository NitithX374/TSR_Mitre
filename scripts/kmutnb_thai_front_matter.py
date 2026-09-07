import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from kmutnb_thai_helpers import (
    setup_page_setup, format_run, add_p, add_h1, add_h2, add_h3, add_h4,
    add_code_block, add_use_case_table, add_data_dict_table,
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_TITLE, FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_TABLE
)

DELIVERABLES_DIR = r"f:\Cybercase Framework\deliverables\thesis_kmutnb_thai"
os.makedirs(DELIVERABLES_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# FRONT MATTER
# ---------------------------------------------------------------------------
def build_front_matter(doc):
    # หน้าปก
    p_seal = doc.add_paragraph()
    p_seal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_seal.paragraph_format.space_before = Pt(36)
    p_seal.paragraph_format.space_after = Pt(24)
    r_seal = p_seal.add_run("[ ตราสัญลักษณ์มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ ]")
    format_run(r_seal, size=Pt(14), italic=True)

    p_title_th = doc.add_paragraph()
    p_title_th.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title_th.paragraph_format.space_before = Pt(30)
    p_title_th.paragraph_format.space_after = Pt(4)
    r_t_th = p_title_th.add_run("ระบบเว็บแอปพลิเคชันจัดการและวิเคราะห์สำนวนคดีอาญาด้วยปัญญาประดิษฐ์")
    format_run(r_t_th, size=FONT_SIZE_TITLE, bold=True)

    p_title_en = doc.add_paragraph()
    p_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title_en.paragraph_format.space_before = Pt(0)
    p_title_en.paragraph_format.space_after = Pt(40)
    r_t_en = p_title_en.add_run("CyberCase Intelligence Framework: An Evidence-Grounded Legal Analysis and Reporting Web Application")
    format_run(r_t_en, size=FONT_SIZE_H1, bold=True)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_before = Pt(40)
    p_author.paragraph_format.space_after = Pt(50)
    r_author = p_author.add_run("นายกฤษกร คำสว่าง")
    format_run(r_author, size=FONT_SIZE_H1, bold=True)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(40)
    p_foot.paragraph_format.space_after = Pt(0)
    p_foot.paragraph_format.line_spacing = 1.15
    r_f1 = p_foot.add_run("ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรปริญญาวิทยาศาสตรบัณฑิต\n")
    format_run(r_f1, size=FONT_SIZE_BODY)
    r_f2 = p_foot.add_run("ภาควิชาวิทยาการคอมพิวเตอร์และสารสนเทศ คณะวิทยาศาสตร์ประยุกต์\n")
    format_run(r_f2, size=FONT_SIZE_BODY)
    r_f3 = p_foot.add_run("มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ\n")
    format_run(r_f3, size=FONT_SIZE_BODY)
    r_f4 = p_foot.add_run("ปีการศึกษา 2568")
    format_run(r_f4, size=FONT_SIZE_BODY)

    doc.add_page_break()

    # หน้าอนุมัติ (ก)
    add_p(doc, "(ก)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, first_indent=0)
    
    p_app_head = doc.add_paragraph()
    p_app_head.paragraph_format.space_before = Pt(10)
    p_app_head.paragraph_format.space_after = Pt(2)
    p_app_head.paragraph_format.line_spacing = 1.15
    r_app_h = p_app_head.add_run("ปริญญานิพนธ์เรื่อง\t: ระบบเว็บแอปพลิเคชันจัดการและวิเคราะห์สำนวนคดีอาญาด้วยปัญญาประดิษฐ์\n\t  CyberCase Intelligence Framework (Backend & Frontend)\nโดย\t\t: นายกฤษกร คำสว่าง\nสาขาวิชา\t\t: วิทยาการคอมพิวเตอร์\nภาควิชา\t\t: วิทยาการคอมพิวเตอร์และสารสนเทศ\nคณะ\t\t: วิทยาศาสตร์ประยุกต์\nอาจารย์ที่ปรึกษา\t: อาจารย์ ดร.สรร รัตนสัญญา\nปีการศึกษา\t: 2568")
    format_run(r_app_h, size=FONT_SIZE_BODY)

    p_app_body = doc.add_paragraph()
    p_app_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_app_body.paragraph_format.space_before = Pt(16)
    p_app_body.paragraph_format.space_after = Pt(20)
    p_app_body.paragraph_format.first_line_indent = Inches(0.5)
    r_ab = p_app_body.add_run("คณะวิทยาศาสตร์ประยุกต์ มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ อนุมัติให้ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรปริญญาวิทยาศาสตรบัณฑิต สาขาวิชาวิทยาการคอมพิวเตอร์")
    format_run(r_ab, size=FONT_SIZE_BODY)

    p_signs = doc.add_paragraph()
    p_signs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_signs.paragraph_format.space_before = Pt(20)
    p_signs.paragraph_format.line_spacing = 1.8
    r_signs = p_signs.add_run(
        "........................................................................... ประธานกรรมการ\n"
        "(ผู้ช่วยศาสตราจารย์ สถิตย์ ประสมพันธ์)\n\n"
        "........................................................................... กรรมการ (อาจารย์ที่ปรึกษา)\n"
        "(อาจารย์ ดร.สรร รัตนสัญญา)\n\n"
        "........................................................................... กรรมการ\n"
        "(ผู้ช่วยศาสตราจารย์ ดร.สุวัจชัย กมลสันติโรจน์)\n\n"
        "........................................................................... กรรมการ\n"
        "(อาจารย์ ดร.ณัฐกิตติ์ จิตรเอื้อตระกูล)\n\n"
        "........................................................................... หัวหน้าภาควิชา\n"
        "(ผู้ช่วยศาสตราจารย์ ดร.อัครา ประโยชน์)\n"
    )
    format_run(r_signs, size=FONT_SIZE_BODY)

    add_p(doc, "ลิขสิทธิ์ของภาควิชาวิทยาการคอมพิวเตอร์และสารสนเทศ คณะวิทยาศาสตร์ประยุกต์\nมหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ\nปีการศึกษา 2568", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, first_indent=0)

    doc.add_page_break()

    # บทคัดย่อ (ข)
    add_p(doc, "(ข)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, first_indent=0)
    add_p(doc, "บทคัดย่อ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=12, bold=True, font_size=FONT_SIZE_TITLE, first_indent=0)
    
    add_p(doc, "ในกระบวนการยุติธรรมทางอาญา พนักงานอัยการมีหน้าที่สำคัญในการตรวจพิจารณาสำนวนการสอบสวนที่พนักงานสอบสวน (ตำรวจ) ส่งมอบมาให้เพื่อมีคำสั่งทางคดี เอกสารในสำนวนมีความหลากหลายและกระจัดกระจาย ประกอบด้วยหนังสือนำส่งสำนวน บันทึกคำให้การพยานและผู้ต้องหา รายงานตรวจสถานที่เกิดเหตุ บัญชีของกลาง ตลอดจนพยานหลักฐานทางอิเล็กทรอนิกส์ในคดีอาญาทั่วไป เช่น สลิปการโอนเงินผ่านโมบายแบงก์กิ้ง และประวัติการสนทนาทางแอปพลิเคชัน LINE ซึ่งเอกสารเหล่านี้มักมีความยาวหลายสิบหรือหลายร้อยหน้า และมีการใช้ถ้อยคำทางเทคนิคปะปนอยู่ แม้ว่าโมเดลภาษาขนาดใหญ่ (Large Language Models: LLMs) จะมีศักยภาพในการสรุปข้อความ แต่การประยุกต์ใช้โมเดลสร้างข้อความทั่วไปในงานคดีอาญากลับสร้างความเสี่ยงร้ายแรง ได้แก่ การเกิดภาพหลอน (Hallucination) การนำข้อกล่าวหาหรือความเห็นสืบสวนมาสรุปเป็นข้อเท็จจริงที่ยุติแล้ว และการสับสนนำเอาข้อมูลภัยคุกคามไซเบอร์ภายนอกมาปะปนกับพฤติการณ์จริงในสำนวน")
    add_p(doc, "เพื่อแก้ไขปัญหาดังกล่าว โครงงานปริญญานิพนธ์นี้จึงได้ออกแบบ พัฒนา และประเมินผลระบบ CyberCase Intelligence Framework โดยเน้นขอบเขตความรับผิดชอบในส่วนระบบส่วนหลัง (Backend API) และระบบส่วนหน้า (Frontend Web Workspace) โดยกำหนดให้ระบบสืบค้นและวิเคราะห์กราฟองค์ความรู้ (GraphRAG Microservice) ทำหน้าที่เป็นบริการภายนอก (Downstream Dependency) ของผู้ร่วมโครงงาน โดยระบบที่พัฒนาขึ้นประกอบด้วยฟังก์ชันหลัก ได้แก่: (1) สถาปัตยกรรม Evidence Trust Boundary ที่กำหนดให้ข้อความและเอกสารที่ผู้ใช้นำเข้าเป็นพยานหลักฐานจริงเพียงแหล่งเดียว และกักกันไม่ให้ข้อมูลเทคนิคภายนอกปะปนกับข้อเท็จจริงในคดี; (2) ระบบวิเคราะห์สำนวนคดีหลักที่สกัดข้อเท็จจริง 5W1H ลำดับเหตุการณ์ตามวันเวลาจริง และจำแนกข้ออ้างออกเป็นสถานะเชิงประจักษ์ (Reported, Inference, Unresolved, Disputed); (3) กลไกการตรวจสอบความถูกต้องของข้อความอ้างอิงตรง (Verbatim Quote Matching) และระบุพิกัดหน้าเอกสารอย่างรัดกุม; (4) นโยบายการถามเพื่อความกระจ่างแบบมีสถานะ (Stateful Clarification Policy) ที่จัดลำดับความสำคัญและตัดลูปการถามซ้ำด้วยระบบ Programmatic Topic Exhaustion; (5) ตัวคัดกรองความเกี่ยวข้องทางไซเบอร์ก่อนการสืบค้น (Pre-Retrieval Applicability Gate) ที่แยกแยะพยานหลักฐานดิจิทัลทั่วไปออกจากพฤติกรรมบุกรุกระบบ; และ (6) ระบบสร้างรายงานสรุปสำนวนคดีอัตโนมัติแบบกำหนดโครงสร้างแน่นอน (Deterministic Report Generation) ในรูปแบบไฟล์ PDF มาตรฐาน")
    add_p(doc, "ผลการดำเนินงานและการทดสอบระบบด้วยชุดทดสอบอัตโนมัติ (351 Automated Pytest Backend Tests และ 150 Frontend Vitest Tests) ร่วมกับการทดสอบด้วยสำนวนการสอบสวนคดีอาญาจริงจากกลุ่มคดีตัวอย่าง ยืนยันว่าระบบที่พัฒนาขึ้นสามารถรักษาความถูกต้องของการอ้างอิงแหล่งที่มาได้สมบูรณ์ ป้องกันการถามวนซ้ำได้ 100% ขจัดปัญหาภาพหลอนในการออกรายงาน และช่วยให้พนักงานอัยการสามารถตรวจสอบความสมบูรณ์ของสำนวนคดีได้อย่างสะดวก รวดเร็ว และแม่นยำยิ่งขึ้น")

    doc.add_page_break()

    # กิตติกรรมประกาศ (ค)
    add_p(doc, "(ค)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, first_indent=0)
    add_p(doc, "กิตติกรรมประกาศ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=12, bold=True, font_size=FONT_SIZE_TITLE, first_indent=0)
    add_p(doc, "ปริญญานิพนธ์ฉบับนี้สำเร็จลุล่วงไปได้ด้วยดีด้วยความกรุณาอย่างยิ่งจาก อาจารย์ ดร.สรร รัตนสัญญา อาจารย์ที่ปรึกษาปริญญานิพนธ์ ที่ได้กรุณาสละเวลาอันมีค่าให้คำปรึกษา คำแนะนำ ตลอดจนถ่ายทอดแนวคิดทางวิชาการและระเบียบวิธีวิจัยที่ถูกต้อง พร้อมทั้งช่วยตรวจสอบและแก้ไขข้อบกพร่องต่าง ๆ ของระบบและเล่มปริญญานิพนธ์มาโดยตลอด ผู้จัดทำขอกราบขอบพระคุณเป็นอย่างสูงไว้ ณ โอกาสนี้")
    add_p(doc, "ขอกราบขอบพระคุณ ผู้ช่วยศาสตราจารย์ สถิตย์ ประสมพันธ์ ประธานกรรมการสอบ, ผู้ช่วยศาสตราจารย์ ดร.สุวัจชัย กมลสันติโรจน์, อาจารย์ ดร.ณัฐกิตติ์ จิตรเอื้อตระกูล กรรมการผู้ทรงคุณวุฒิ และผู้ช่วยศาสตราจารย์ ดร.อัครา ประโยชน์ หัวหน้าภาควิชาวิทยาการคอมพิวเตอร์และสารสนเทศ ที่ได้กรุณาให้ข้อเสนอแนะ ข้อคิดเห็นเชิงวิชาการ และคำแนะนำอันทรงคุณค่ายิ่งในการพัฒนาระบบและการตรวจทานเอกสารฉบับนี้")
    add_p(doc, "ขอขอบคุณคณาจารย์ประจำภาควิชาวิทยาการคอมพิวเตอร์และสารสนเทศทุกท่าน ที่ได้ประสิทธิ์ประสาทวิชาความรู้ ทักษะทางด้านการเขียนโปรแกรม วิศวกรรมซอฟต์แวร์ และปัญญาประดิษฐ์ ตลอดระยะเวลาการศึกษา และขอขอบคุณเพื่อนร่วมโครงงานที่ร่วมกันพัฒนาในส่วนของการสืบค้นองค์ความรู้ภายนอกจนทำให้ระบบสำเร็จสมบูรณ์")
    add_p(doc, "ท้ายที่สุดนี้ ผู้จัดทำขอกราบขอบพระคุณ บิดา มารดา และสมาชิกในครอบครัวทุกท่าน ที่ให้การสนับสนุนทั้งกำลังใจ กำลังทรัพย์ และความอบอุ่นเสมอมา ตลอดจนเพื่อน ๆ ทุกคนที่คอยให้ความช่วยเหลือ ให้คำปรึกษา และร่วมฟันฝ่าอุปสรรคต่าง ๆ จนกระทั่งปริญญานิพนธ์ฉบับนี้สำเร็จลุล่วงได้ด้วยดี")
    add_p(doc, "นายกฤษกร คำสว่าง\nมีนาคม 2569", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=20, first_indent=0)

    doc.add_page_break()

    # สารบัญ (ง)
    add_p(doc, "(ง)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, first_indent=0)
    add_p(doc, "สารบัญ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=14, bold=True, font_size=FONT_SIZE_TITLE, first_indent=0)
    
    toc_items = [
        ("บทคัดย่อ", "ข"),
        ("กิตติกรรมประกาศ", "ค"),
        ("สารบัญ", "ง"),
        ("สารบัญตาราง", "ฉ"),
        ("สารบัญภาพ", "ช"),
        ("บทที่ 1 บทนำ", "1"),
        ("       1.1 ความเป็นมาและความสำคัญของโครงงาน", "1"),
        ("       1.2 วัตถุประสงค์ของโครงงาน", "3"),
        ("       1.3 ขอบเขตของโครงงาน", "4"),
        ("       1.4 ประโยชน์ที่คาดว่าจะได้รับ", "6"),
        ("บทที่ 2 ทฤษฎีและงานวิจัยที่เกี่ยวข้อง", "7"),
        ("       2.1 บริบทงานกระบวนการยุติธรรมและเอกสารการสอบสวน", "7"),
        ("       2.2 พยานหลักฐานดิจิทัลและกรอบความรู้ภัยคุกคามไซเบอร์", "9"),
        ("       2.3 โมเดลภาษาขนาดใหญ่และการค้นคืนข้อมูล", "11"),
        ("       2.4 การอ้างอิงแหล่งที่มาและการพิสูจน์ข้อเท็จจริง", "13"),
        ("       2.5 การถามเพื่อความกระจ่างในระบบสนทนา", "15"),
        ("       2.6 เครื่องมือและสถาปัตยกรรมซอฟต์แวร์ที่ใช้ในการพัฒนา", "17"),
        ("บทที่ 3 ขั้นตอนและวิธีการดำเนินงาน", "21"),
        ("       3.1 System Architecture", "21"),
        ("       3.2 Flowchart แสดงการไหลของข้อมูลในระบบ", "23"),
        ("       3.3 Use Case Diagram ฟังก์ชันการทำงานของระบบ", "27"),
        ("       3.4 Use Case Description", "29"),
        ("       3.5 Entity-Relationship Diagram (ER Diagram)", "45"),
        ("       3.6 Data Dictionary", "46"),
        ("บทที่ 4 การพัฒนาระบบ", "52"),
        ("       4.1 เครื่องมือและสภาพแวดล้อมที่ใช้ในการพัฒนาระบบ", "52"),
        ("       4.2 การพัฒนาระบบของโครงงานฝั่ง Backend", "53"),
        ("       4.3 การพัฒนาระบบของโครงงานฝั่ง Frontend", "68"),
        ("บทที่ 5 ผลการดำเนินโครงการและการแสดงผล", "76"),
        ("       5.1 ผลการดำเนินงานภาพรวม", "76"),
        ("       5.2 การแสดงผลของ Web Application", "77"),
        ("       5.3 ผลการทดสอบความถูกต้องและประสิทธิภาพของระบบ", "92"),
        ("บทที่ 6 บทสรุป ปัญหา และแนวทางในการพัฒนาต่อ", "98"),
        ("       6.1 สรุปผลการดำเนินงาน", "98"),
        ("       6.2 ปัญหาที่พบระหว่างการดำเนินงาน และแนวทางการแก้ปัญหา", "100"),
        ("       6.3 แนวทางในการพัฒนาต่อ", "103"),
        ("บรรณานุกรม", "105")
    ]
    
    p_toc_head = doc.add_paragraph()
    p_toc_head.paragraph_format.space_before = Pt(0)
    p_toc_head.paragraph_format.space_after = Pt(4)
    r_th1 = p_toc_head.add_run("เรื่อง\t\t\t\t\t\t\t\t\t\tหน้า")
    format_run(r_th1, size=FONT_SIZE_BODY, bold=True)
    
    for title, page in toc_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_before = Pt(1)
        p_item.paragraph_format.space_after = Pt(1)
        p_item.paragraph_format.line_spacing = 1.1
        is_b = "บทที่" in title or title in ["บทคัดย่อ", "กิตติกรรมประกาศ", "สารบัญ", "บรรณานุกรม"]
        r = p_item.add_run(f"{title}")
        format_run(r, size=FONT_SIZE_BODY, bold=is_b)
        # Tab stop
        r_tab = p_item.add_run(f"\t\t\t\t\t\t\t\t\t\t{page}")
        format_run(r_tab, size=FONT_SIZE_BODY, bold=is_b)

    doc.add_page_break()

    # สารบัญตาราง (ฉ)
    add_p(doc, "(ฉ)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, first_indent=0)
    add_p(doc, "สารบัญตาราง", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=14, bold=True, font_size=FONT_SIZE_TITLE, first_indent=0)
    
    tables_list = [
        ("ตารางที่ 3-1 ตาราง Use Case Description : Create New Case Workspace", "29"),
        ("ตารางที่ 3-2 ตาราง Use Case Description : Upload & Extract Case Document", "30"),
        ("ตารางที่ 3-3 ตาราง Use Case Description : Preview & Confirm Narrative", "31"),
        ("ตารางที่ 3-4 ตาราง Use Case Description : Execute Main Case Analysis", "32"),
        ("ตารางที่ 3-5 ตาราง Use Case Description : Classify Technical Applicability", "33"),
        ("ตารางที่ 3-6 ตาราง Use Case Description : Request Downstream Technical Intelligence", "34"),
        ("ตารางที่ 3-7 ตาราง Use Case Description : Extract 5W1H Overview & Chronology", "35"),
        ("ตารางที่ 3-8 ตาราง Use Case Description : Categorize Epistemic Claims", "36"),
        ("ตารางที่ 3-9 ตาราง Use Case Description : Validate Verbatim Source Quotes", "37"),
        ("ตารางที่ 3-10 ตาราง Use Case Description : Detect Evidentiary Gaps", "38"),
        ("ตารางที่ 3-11 ตาราง Use Case Description : Select & Formulate Clarification Question", "39"),
        ("ตารางที่ 3-12 ตาราง Use Case Description : Exhaust Clarification Topic", "40"),
        ("ตารางที่ 3-13 ตาราง Use Case Description : Inspect Source Passage in Drawer", "41"),
        ("ตารางที่ 3-14 ตาราง Use Case Description : Ask Contextual Question", "42"),
        ("ตารางที่ 3-15 ตาราง Use Case Description : Generate Deterministic Forensic Report", "43"),
        ("ตารางที่ 3-16 ตาราง Use Case Description : Retry Interrupted Processing Run", "44"),
        ("ตารางที่ 3-17 Data Dictionary : chat_threads", "47"),
        ("ตารางที่ 3-18 Data Dictionary : chat_messages", "48"),
        ("ตารางที่ 3-19 Data Dictionary : chat_runs", "49"),
        ("ตารางที่ 3-20 Data Dictionary : rag_contexts", "50"),
        ("ตารางที่ 3-21 Data Dictionary : chat_reports", "51"),
        ("ตารางที่ 5-1 สรุปผลการทดสอบระบบส่วนหลัง (Backend Pytest Receipts)", "93"),
        ("ตารางที่ 5-2 สรุปผลการทดสอบระบบส่วนหน้า (Frontend Vitest Receipts)", "94"),
        ("ตารางที่ 5-3 สรุปผลการประเมินความแม่นยำบนสำนวนการสอบสวนคดีจริง 3 รูปแบบ", "96")
    ]
    
    p_tab_head = doc.add_paragraph()
    p_tab_head.paragraph_format.space_after = Pt(4)
    r_th2 = p_tab_head.add_run("ตารางที่\t\t\t\t\t\t\t\t\tหน้า")
    format_run(r_th2, size=FONT_SIZE_BODY, bold=True)
    
    for t_title, t_page in tables_list:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_before = Pt(1)
        p_item.paragraph_format.space_after = Pt(1)
        p_item.paragraph_format.line_spacing = 1.1
        r = p_item.add_run(f"{t_title}")
        format_run(r, size=FONT_SIZE_BODY)
        r_tab = p_item.add_run(f"\t\t\t\t\t\t\t\t\t{t_page}")
        format_run(r_tab, size=FONT_SIZE_BODY)

    doc.add_page_break()

    # สารบัญภาพ (ช)
    add_p(doc, "(ช)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, first_indent=0)
    add_p(doc, "สารบัญภาพ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=14, bold=True, font_size=FONT_SIZE_TITLE, first_indent=0)
    
    figures_list = [
        ("ภาพที่ 3-1 System Architecture สถาปัตยกรรมระบบรวม CyberCase", "22"),
        ("ภาพที่ 3-2 Flowchart แสดงขั้นตอนการเตรียมและสกัดเอกสารสำนวนคดี (Intake)", "23"),
        ("ภาพที่ 3-3 Flowchart แสดงขั้นตอนการประมวลผลวิเคราะห์หลัก (Pipeline Execution)", "24"),
        ("ภาพที่ 3-4 Flowchart แสดงการคัดกรองความเกี่ยวข้องทางไซเบอร์ (Applicability Gate)", "25"),
        ("ภาพที่ 3-5 Flowchart แสดงขั้นตอนการถามตอบความชัดเจนและตัดลูป (Clarification)", "26"),
        ("ภาพที่ 3-6 Flowchart แสดงขั้นตอนการสร้างรายงานสรุปสำนวนคดี (Report Generation)", "27"),
        ("ภาพที่ 3-7 Use Case Diagram ฟังก์ชันการทำงานของระบบ CyberCase", "28"),
        ("ภาพที่ 3-8 Entity-Relationship Diagram (ER Diagram) โครงสร้างฐานข้อมูล", "46"),
        ("ภาพที่ 4-1 Source code การสกัดและแปลงข้อความเอกสาร (preview.py)", "54"),
        ("ภาพที่ 4-2 Source code ตัวคัดกรองความเกี่ยวข้องทางไซเบอร์ (applicability_gate.py)", "56"),
        ("ภาพที่ 4-3 Source code โครงสร้าง Pydantic Trace V3 และ Claim Enums (contracts.py)", "58"),
        ("ภาพที่ 4-4 Source code การถอดรหัสและการตรวจสอบโครงสร้าง Trace (response_decoder.py)", "60"),
        ("ภาพที่ 4-5 Source code อัลกอริทึมการตรวจสอบข้อความอ้างอิงตรง (source_citations.py)", "62"),
        ("ภาพที่ 4-6 Source code นโยบายการคัดเลือกประเด็นคำถามและการตัดลูป (decision.py)", "64"),
        ("ภาพที่ 4-7 Source code การจัดการสัญญาเช่า Worker และการกู้คืนงานขัดข้อง (locks.py)", "66"),
        ("ภาพที่ 4-8 Source code ตัวแปลงมุมมองรายงานและการเรนเดอร์ PDF (report_builder.py)", "67"),
        ("ภาพที่ 4-9 Source code การจัดการแคชและการ Polling แบบวงรอบเดียว (chat-polling.ts)", "69"),
        ("ภาพที่ 4-10 Source code การแยกสถานะแบบร่างและการส่งข้อความซ้ำ (submission.ts)", "71"),
        ("ภาพที่ 4-11 Source code คอมโพเนนต์แสดงภาพรวมและข้ออ้างเชิงประจักษ์ (Findings.tsx)", "73"),
        ("ภาพที่ 4-12 Source code คอมโพเนนต์ลิ้นชักตรวจสอบเอกสารและการเน้นข้อความ (Drawer.tsx)", "75"),
        ("ภาพที่ 5-1 หน้าจอ Case Intake Preparation และการอัปโหลดไฟล์เอกสาร", "78"),
        ("ภาพที่ 5-2 หน้าต่างแสดงตัวอย่างข้อความสกัดก่อนการกดยืนยัน (Document Preview)", "80"),
        ("ภาพที่ 5-3 หน้าจอแสดงสถานะการประมวลผลเบื้องหลัง (Background Run Stepper)", "81"),
        ("ภาพที่ 5-4 หน้าจอภาพรวมคดี (Case Overview) แสดงสรุป 5W1H และลำดับเหตุการณ์", "83"),
        ("ภาพที่ 5-5 หน้าจอข้อกล่าวหาและข้อเท็จจริงจำแนกตามสถานะเชิงประจักษ์ (Grounded Findings)", "85"),
        ("ภาพที่ 5-6 หน้าต่างลิ้นชักตรวจสอบหลักฐานต้นทางและการเน้นข้อความตรง (Citation Drawer)", "87"),
        ("ภาพที่ 5-7 หน้าต่างแสดงประเด็นสงสัยและคำถามขอความกระจ่าง (Clarification Card)", "88"),
        ("ภาพที่ 5-8 หน้าต่างคลังข้อมูลเทคนิคภายนอก MITRE ATT&CK (Technical Context Drawer)", "89"),
        ("ภาพที่ 5-9 หน้าจอแสดงรายงานสรุปสำนวนคดีแบบตอบสนอง (Interactive Report View)", "90"),
        ("ภาพที่ 5-10 ตัวอย่างเอกสารรายงานสรุปสำนวนคดีรูปแบบไฟล์ PDF ทางการ", "91"),
        ("ภาพที่ 5-11 หน้าต่างแจ้งเตือนข้อผิดพลาดและการกู้คืนงานประมวลผล (Run Recovery Modal)", "92")
    ]
    
    p_fig_head = doc.add_paragraph()
    p_fig_head.paragraph_format.space_after = Pt(4)
    r_th3 = p_fig_head.add_run("ภาพที่\t\t\t\t\t\t\t\t\tหน้า")
    format_run(r_th3, size=FONT_SIZE_BODY, bold=True)
    
    for f_title, f_page in figures_list:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_before = Pt(1)
        p_item.paragraph_format.space_after = Pt(1)
        p_item.paragraph_format.line_spacing = 1.1
        r = p_item.add_run(f"{f_title}")
        format_run(r, size=FONT_SIZE_BODY)
        r_tab = p_item.add_run(f"\t\t\t\t\t\t\t\t\t{f_page}")
        format_run(r_tab, size=FONT_SIZE_BODY)

    doc.add_page_break()

print("Front matter builder ready.")
