import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

FONT_NAME = "TH SarabunPSK"
FONT_SIZE_BODY = Pt(16)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(16)
FONT_SIZE_TITLE = Pt(20)
FONT_SIZE_CODE = Pt(13)
FONT_SIZE_TABLE = Pt(14)

def setup_page_setup(section):
    section.top_margin = Inches(1.5)
    section.left_margin = Inches(1.5)
    section.bottom_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)

def format_run(run, font_name=FONT_NAME, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}" w:eastAsia="{font_name}"/>')
    rPr.append(rFonts)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="B0BEC5", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line_spacing=1.15, first_indent=0.5, bold=False, italic=False, font_size=FONT_SIZE_BODY, font_name=FONT_NAME, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if text:
        run = p.add_run(text)
        format_run(run, font_name=font_name, size=font_size, bold=bold, italic=italic, color=color)
    return p

def add_h1(doc, chapter_num, chapter_title):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(18)
    p1.paragraph_format.space_after = Pt(4)
    run1 = p1.add_run(chapter_num)
    format_run(run1, size=FONT_SIZE_TITLE, bold=True)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(14)
    run2 = p2.add_run(chapter_title)
    format_run(run2, size=FONT_SIZE_TITLE, bold=True)

def add_h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(title)
    format_run(run, size=FONT_SIZE_H1, bold=True)
    return p

def add_h3(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(title)
    format_run(run, size=FONT_SIZE_H2, bold=True)
    return p

def add_h4(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(title)
    format_run(run, size=FONT_SIZE_H3, bold=True)
    return p

def add_code_block(doc, title, code_text, explanation_text):
    # Caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(4)
    r_cap = cp.add_run(title)
    format_run(r_cap, size=Pt(14), bold=True)
    
    # Table box for code
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
        f'  <w:left w:val="single" w:sz="20" w:space="0" w:color="0284C7"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    format_run(run, font_name="Consolas", size=FONT_SIZE_CODE, color=RGBColor(30, 41, 59))
    
    # Explanation
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ep.paragraph_format.first_line_indent = Inches(0.5)
    ep.paragraph_format.space_before = Pt(6)
    ep.paragraph_format.space_after = Pt(8)
    r_exp = ep.add_run(explanation_text)
    format_run(r_exp, size=FONT_SIZE_BODY)

def add_use_case_table(doc, table_num, uc_num, uc_name, brief_desc, actors, pre_cond, post_cond, main_flow_actor, main_flow_system, alt_flow, explanation):
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.space_before = Pt(10)
    cp.paragraph_format.space_after = Pt(4)
    run_cap = cp.add_run(f"ตารางที่ {table_num} Use Case Description : {uc_name}")
    format_run(run_cap, size=FONT_SIZE_H2, bold=True)
    
    tbl = doc.add_table(rows=7, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="94A3B8")
    
    headers = [
        ("Use Case Number", str(uc_num)),
        ("Use Case Name", uc_name),
        ("Brief Description", brief_desc),
        ("Actors", actors),
        ("Pre-condition", pre_cond),
        ("Post-Condition", post_cond),
    ]
    
    for r_idx, (k, v) in enumerate(headers):
        c0 = tbl.cell(r_idx, 0)
        c1 = tbl.cell(r_idx, 1)
        c0.width = Inches(2.0)
        c1.width = Inches(4.5)
        set_cell_background(c0, "F1F5F9")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(k)
        format_run(r0, size=FONT_SIZE_TABLE, bold=True)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(v)
        format_run(r1, size=FONT_SIZE_TABLE)
        
    # Main Flow row (Row 6)
    row_mf = tbl.rows[6]
    c0 = row_mf.cells[0]
    c1 = row_mf.cells[1]
    c0.width = Inches(2.0)
    c1.width = Inches(4.5)
    set_cell_background(c0, "F1F5F9")
    set_cell_margins(c0, 60, 60, 100, 100)
    set_cell_margins(c1, 60, 60, 100, 100)
    
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("Main Flow")
    format_run(r0, size=FONT_SIZE_TABLE, bold=True)
    
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1_a = p1.add_run("Actor:\n")
    format_run(r1_a, size=FONT_SIZE_TABLE, bold=True)
    r1_at = p1.add_run(f"{main_flow_actor}\n\n")
    format_run(r1_at, size=FONT_SIZE_TABLE)
    r1_s = p1.add_run("System:\n")
    format_run(r1_s, size=FONT_SIZE_TABLE, bold=True)
    r1_st = p1.add_run(main_flow_system)
    format_run(r1_st, size=FONT_SIZE_TABLE)
    
    # Alternate Flow row
    row_af = tbl.add_row()
    c0 = row_af.cells[0]
    c1 = row_af.cells[1]
    c0.width = Inches(2.0)
    c1.width = Inches(4.5)
    set_cell_background(c0, "F1F5F9")
    set_cell_margins(c0, 60, 60, 100, 100)
    set_cell_margins(c1, 60, 60, 100, 100)
    
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("Alternate / Exceptional Flows")
    format_run(r0, size=FONT_SIZE_TABLE, bold=True)
    
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(alt_flow)
    format_run(r1, size=FONT_SIZE_TABLE)
    
    # Explanation
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ep.paragraph_format.first_line_indent = Inches(0.5)
    ep.paragraph_format.space_before = Pt(4)
    ep.paragraph_format.space_after = Pt(8)
    r_exp = ep.add_run(f"อธิบายตารางที่ {table_num} คือ {explanation}")
    format_run(r_exp, size=FONT_SIZE_BODY)

def add_data_dict_table(doc, table_num, table_name, table_desc, rows_data):
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.space_before = Pt(10)
    cp.paragraph_format.space_after = Pt(2)
    run_cap = cp.add_run(f"ตารางที่ {table_num} Data Dictionary : {table_name}")
    format_run(run_cap, size=FONT_SIZE_H2, bold=True)
    
    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(0)
    dp.paragraph_format.space_after = Pt(4)
    dp.paragraph_format.first_line_indent = Inches(0.2)
    run_desc = dp.add_run(f"{table_name} จัดเก็บ{table_desc}")
    format_run(run_desc, size=FONT_SIZE_TABLE, italic=True)
    
    col_names = ["Key", "Column Name", "Data Type", "Nullable", "Description", "Reference"]
    col_widths = [Inches(0.6), Inches(1.5), Inches(1.2), Inches(1.0), Inches(1.8), Inches(1.1)]
    
    tbl = doc.add_table(rows=len(rows_data)+1, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="94A3B8")
    
    # Header
    for c_idx, name in enumerate(col_names):
        cell = tbl.cell(0, c_idx)
        cell.width = col_widths[c_idx]
        set_cell_background(cell, "E2E8F0")
        set_cell_margins(cell, 60, 60, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(name)
        format_run(r, size=FONT_SIZE_TABLE, bold=True)
        
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx+1, c_idx)
            cell.width = col_widths[c_idx]
            set_cell_margins(cell, 40, 40, 80, 80)
            if c_idx == 0:
                set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            if c_idx in [0, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            format_run(r, size=FONT_SIZE_TABLE)
            
    post_p = doc.add_paragraph()
    post_p.paragraph_format.space_after = Pt(6)

print("Base helper loaded.")
