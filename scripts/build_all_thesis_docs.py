import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from convert_thesis_to_docx import (
    convert_markdown_to_docx, setup_page_setup, add_formatted_text, format_run,
    FONT_SIZE_TITLE, FONT_SIZE_BODY, FONT_NAME
)

SOURCE_DIR = r"f:\Cybercase Framework\deliverables\thesis_backend_frontend\source"
DELIVERABLES_DIR = r"f:\Cybercase Framework\deliverables\thesis_backend_frontend"
ROOT_DIR = r"f:\Cybercase Framework"

CHAPTERS = [
    ("chapter_i.md", "Chapter_I_Introduction.docx", "Chapter I: Introduction"),
    ("chapter_ii.md", "Chapter_II_Literature_Review.docx", "Chapter II: Literature Review"),
    ("chapter_iii.md", "Chapter_III_Research_Methodology.docx", "Chapter III: Research Methodology"),
    ("chapter_iv.md", "Chapter_IV_System_Development.docx", "Chapter IV: System Development"),
    ("chapter_v.md", "Chapter_V_Verification_and_Evaluation.docx", "Chapter V: System Verification and Evaluation"),
    ("chapter_vi.md", "Chapter_VI_Conclusion_and_Future_Work.docx", "Chapter VI: Conclusion and Future Work")
]

def build_individual_chapters():
    print("--- Generating Individual Chapter Word Documents ---")
    for md_file, docx_file, title in CHAPTERS:
        md_path = os.path.join(SOURCE_DIR, md_file)
        docx_path = os.path.join(DELIVERABLES_DIR, docx_file)
        if not os.path.exists(md_path):
            print(f"Warning: {md_path} not found.")
            continue
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        doc = docx.Document()
        setup_page_setup(doc.sections[0])
        convert_markdown_to_docx(content, doc)
        doc.save(docx_path)
        print(f"Generated: {docx_path} ({os.path.getsize(docx_path)} bytes)")

def build_consolidated_thesis():
    print("\n--- Generating Consolidated Full KMUTNB Thesis Monograph ---")
    template_path = os.path.join(ROOT_DIR, "Cybercase_Thesis_KMUTNB_Full.docx")
    backup_path = os.path.join(ROOT_DIR, "Cybercase_Thesis_KMUTNB_Full.backup.docx")
    
    if os.path.exists(template_path) and not os.path.exists(backup_path):
        shutil.copyfile(template_path, backup_path)
        print(f"Saved original template backup to: {backup_path}")
        
    full_doc = docx.Document()
    setup_page_setup(full_doc.sections[0])
    
    # Check if we can extract front matter from template
    if os.path.exists(template_path):
        src_template = docx.Document(template_path)
        print("Extracting official front matter (Title, Approval, Abstract, Acknowledgements, TOC)...")
        
        # Paragraphs 0 to 19 contain the front matter
        for i, p in enumerate(src_template.paragraphs):
            if "Chapter 1" in p.text or "CHAPTER 1" in p.text or "Chapter I" in p.text:
                break
            
            # Update abstract if this is the abstract paragraph (index 8)
            text_to_add = p.text
            if "Modern cybersecurity operations face severe bottlenecks" in text_to_add:
                text_to_add = (
                    "In criminal justice and forensic review, public prosecutors face significant cognitive bottlenecks when reviewing "
                    "complex police investigation dossiers (รายงานการสอบสวน). These case files combine diverse information: formal procedural "
                    "transmittal letters, witness depositions, physical damage claims, financial transfer slips, and mobile communication records. "
                    "While Large Language Models (LLMs) offer strong narrative processing capabilities, general-purpose systems frequently fail in "
                    "legal contexts by conflating witness allegations with established facts, hallucinating source citations, and confusing "
                    "external cybersecurity taxonomies with case facts.\n\n"
                    "To resolve these challenges, this thesis presents the full-stack design, engineering, and empirical verification of CyberCase "
                    "Intelligence Framework, an evidence-grounded legal analysis and reporting platform. The candidate's primary development "
                    "responsibility encompasses the complete Frontend and Backend system tiers, treating downstream graph-retrieval as an external "
                    "partner microservice. The research develops: (1) an Evidence Trust Model enforcing strict boundaries between authoritative "
                    "case evidence and analytical context; (2) a multi-stage analysis pipeline extracting 5W1H overviews, temporal event chronologies, "
                    "and discrete claims tagged with explicit epistemic certainty statuses (reported, inference, unresolved, disputed); (3) a "
                    "fail-closed source provenance engine validating verbatim quotation containment and conservative document page coordinates; "
                    "(4) a stateful clarification engine with priority ranking and programmatic topic exhaustion preventing conversational loops; "
                    "(5) an explicit pre-retrieval applicability gate distinguishing ordinary digital artifacts from cyber-adversarial techniques; "
                    "and (6) a deterministic, template-first report assembly engine that compiles audit-ready forensic PDF dossiers without secondary "
                    "generative drift.\n\n"
                    "Empirical verification across automated Pytest test suites (351 backend and 150 frontend tests) and domain case typologies from "
                    "authentic Thai police files demonstrates that the developed full-stack platform maintains complete provenance auditability, "
                    "enforces 100% loop prevention, eliminates external knowledge contamination, and provides an auditable cognitive workspace for "
                    "legal case review."
                )
            
            new_p = full_doc.add_paragraph()
            new_p.alignment = p.alignment
            new_p.paragraph_format.space_before = p.paragraph_format.space_before
            new_p.paragraph_format.space_after = p.paragraph_format.space_after
            new_p.paragraph_format.line_spacing = p.paragraph_format.line_spacing
            
            # Determine bold / size based on text
            is_bold = any(r.bold for r in p.runs) or p.text.isupper() or "ABSTRACT" in p.text or "APPROVAL" in p.text or "TABLE OF CONTENTS" in p.text
            f_size = FONT_SIZE_TITLE if (i == 0 or "ABSTRACT" in p.text or "APPROVAL" in p.text or "TABLE" in p.text) else FONT_SIZE_BODY
            
            if i in [0, 1, 2, 4, 7, 11, 15, 17, 18]:
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = new_p.add_run(text_to_add)
            format_run(run, size=f_size, bold=is_bold)
            
            # Page break after major front-matter sections
            if "Academic Year 2026" in p.text or "LIST OF FIGURES" in p.text or "FOR THE DEGREE OF" in p.text:
                full_doc.add_page_break()

    # Append all 6 chapters with page breaks
    for md_file, _, title in CHAPTERS:
        md_path = os.path.join(SOURCE_DIR, md_file)
        if not os.path.exists(md_path):
            continue
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        full_doc.add_page_break()
        print(f"Appending {title}...")
        convert_markdown_to_docx(content, full_doc)

    out_deliverable = os.path.join(DELIVERABLES_DIR, "Cybercase_Thesis_KMUTNB_Full_Backend_Frontend.docx")
    full_doc.save(out_deliverable)
    print(f"Saved consolidated deliverable: {out_deliverable} ({os.path.getsize(out_deliverable)} bytes)")
    
    out_root = os.path.join(ROOT_DIR, "Cybercase_Thesis_KMUTNB_Full.docx")
    full_doc.save(out_root)
    print(f"Updated root thesis document: {out_root} ({os.path.getsize(out_root)} bytes)")

if __name__ == "__main__":
    build_individual_chapters()
    build_consolidated_thesis()
    print("\nAll Word documents generated successfully!")
