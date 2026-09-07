import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from kmutnb_thai_helpers import (
    add_p, add_h1, add_h2, add_h3, add_h4,
    add_use_case_table, add_data_dict_table,
    FONT_SIZE_BODY, FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_H3
)

def build_chapter_3(doc):
    add_h1(doc, "บทที่ 3", "ขั้นตอนและวิธีการดำเนินงาน")
    
    add_p(doc, "ในบทนี้จะกล่าวถึงโครงสร้างการออกแบบสถาปัตยกรรมระบบ (System Architecture) ของระบบ CyberCase Intelligence Framework การไหลของข้อมูลในระบบผ่านแผนผัง Flowchart แผนภาพฟังก์ชันการทำงานของระบบ (Use Case Diagram) รายละเอียดการทำงานของแต่ละ Use Case อย่างละเอียด (Use Case Descriptions) ตลอดจนแผนภาพความสัมพันธ์ของข้อมูล (ER Diagram) และพจนานุกรมข้อมูล (Data Dictionary) ที่จัดเก็บในฐานข้อมูล PostgreSQL")
    
    add_h2(doc, "3.1 System Architecture")
    add_p(doc, "สถาปัตยกรรมของระบบ CyberCase Intelligence Framework ได้รับการออกแบบตามหลักการแยกหน้าที่ความรับผิดชอบ (Separation of Concerns) และการควบคุมขอบเขตความน่าเชื่อถือของพยานหลักฐาน (Evidence Trust Boundary) โดยแบ่งออกเป็น 4 ชั้นหลัก ได้แก่ Frontend Client Workspace, Backend Application API, Relational Persistence Layer และ Downstream Knowledge Service ดังแสดงในภาพที่ 3-1")
    
    # ภาพสถาปัตยกรรม
    p_fig_arch = doc.add_paragraph()
    p_fig_arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig_arch.paragraph_format.space_before = Pt(8)
    p_fig_arch.paragraph_format.space_after = Pt(4)
    r_fa = p_fig_arch.add_run("ภาพที่ 3-1 System Architecture สถาปัตยกรรมระบบรวม CyberCase")
    r_fa.bold = True
    
    add_p(doc, "จากภาพที่ 3-1 จะเห็นการแบ่งสถาปัตยกรรมระบบออกเป็น 4 ส่วนสำคัญ:")
    add_p(doc, "1. Frontend Web Workspace (Next.js 16 / React 19): เป็นส่วนติดต่อผู้ใช้สำหรับพนักงานอัยการ ทำหน้าที่จัดการหน้าจอการเตรียมเอกสารสำนวนคดี (Intake Workspace), หน้าจอแสดงภาพรวม 5W1H และข้ออ้างเชิงประจักษ์ (Case Overview), หน้าจอลิ้นชักตรวจสอบหลักฐานต้นทาง (Citation Drawer), การ์ดตอบสนองคำถามเพื่อความกระจ่าง และหน้าจอรายงานสรุปสำนวนคดี โดยมี TanStack Query เป็นตัวควบคุมแคชและ Polling สถานะจากเซิร์ฟเวอร์", first_indent=0.75)
    add_p(doc, "2. Backend Application Services (FastAPI Async): ทำหน้าที่เป็นแกนกลางในการประมวลผล ประกอบด้วย Case Analysis Service (สร้าง Prompt, สกัด 5W1H, ถอดรหัสโครงสร้าง Pydantic), Validation Engine (ตรวจการอ้างอิงตรงและพิกัดหน้า), Follow-up Service (วิเคราะห์ช่องว่างคดีและตัดลูป), Applicability Gate (คัดกรองความเกี่ยวข้องทางไซเบอร์), Leased Worker (ประมวลผลเบื้องหลังแบบเช่าเวลา) และ Report Assembly Service (ประกอบรายงาน PDF)", first_indent=0.75)
    add_p(doc, "3. Persistence Layer (PostgreSQL): จัดเก็บข้อมูลสถานะของคดีอย่างถาวรและมีความสัมพันธ์กันอย่างเข้มงวด ได้แก่ ข้อมูลห้องสำนวน (chat_threads), ประวัติข้อความหลักฐานและข้อความระบบ (chat_messages), ข้อมูลการประมวลผลเบื้องหลังพร้อมสัญญาเช่า (chat_runs), ข้อมูลบริบทที่สืบค้นได้จาก RAG (rag_contexts) และข้อมูลรายงานสรุปสำนวนคดี (chat_reports)", first_indent=0.75)
    add_p(doc, "4. Downstream Knowledge Service (GraphRAG): เป็นบริการปลายน้ำของผู้ร่วมโครงงานที่รวบรวมฐานข้อมูล STIX 2.1 / MITRE ATT&CK โดยระบบส่วนหลังจะเรียกใช้งานผ่าน REST API เมื่อพยานหลักฐานในคดีบ่งชี้ถึงพฤติการณ์ทางไซเบอร์จริงเท่านั้น", first_indent=0.75)

    add_h2(doc, "3.2 Flowchart แสดงการไหลของข้อมูลในระบบ")
    add_p(doc, "การทำงานของระบบ CyberCase มีเส้นทางการไหลของข้อมูลที่แน่นอนและเป็นระบบ โดยสามารถอธิบายผ่านแผนผังการไหลของข้อมูล (Flowchart) 5 แผนผังหลัก ดังนี้:")
    
    add_h3(doc, "3.2.1 การนำเข้าและเตรียมเอกสารสำนวนคดี (Intake Flowchart)")
    add_p(doc, "ภาพที่ 3-2 แสดงขั้นตอนการไหลของข้อมูลในส่วนของการนำเข้าเอกสารสำนวนคดี:")
    add_p(doc, "1. ผู้ใช้ (พนักงานอัยการ) ทำการอัปโหลดไฟล์เอกสารสำนวนคดี (PDF หรือ DOCX) เข้าสู่ระบบ", first_indent=0.75)
    add_p(doc, "2. ระบบส่วนหลังเรียกใช้โมดูล Document Ingestion ทำการสกัดข้อความ จัดเรียงย่อหน้า และประเมินค่าความเชื่อมั่นของตัวอักษร (OCR Confidence)", first_indent=0.75)
    add_p(doc, "3. ระบบส่งข้อความที่สกัดได้กลับมายังส่วนหน้าเพื่อแสดงผลในหน้าต่าง Document Preview Modal โดยยังไม่มีการบันทึกเป็นพยานหลักฐานคดีและยังไม่เรียกใช้โมเดลภาษา", first_indent=0.75)
    add_p(doc, "4. ผู้ใช้ทำการอ่านตรวจทาน แก้ไขข้อความที่อาจสกัดผิดพลาด และกดปุ่มยืนยัน (Confirm Narrative)", first_indent=0.75)
    add_p(doc, "5. ระบบจึงทำการบันทึกข้อความดังกล่าวลงในฐานข้อมูลเป็น Authoritative Evidence Message (m_1) ที่มีรหัสประจำตัวและแฮช SHA-256 ป้องกันการแก้ไข พร้อมสร้าง Background Run เพื่อเริ่มการวิเคราะห์", first_indent=0.75)

    add_h3(doc, "3.2.2 การประมวลผลวิเคราะห์สำนวนคดีหลัก (Pipeline Execution Flowchart)")
    add_p(doc, "ภาพที่ 3-3 แสดงขั้นตอนการประมวลผลของ Pipeline Execution ภายในระบบส่วนหลัง:")
    add_p(doc, "1. Leased Worker ทำการเคลมสิทธิ์ในการรันงาน (Claim Lease) จากฐานข้อมูลด้วยระยะเวลา 6 นาที และเริ่มส่งสัญญาณ Heartbeat ทุก 30 วินาที", first_indent=0.75)
    add_p(doc, "2. ระบบเรียกใช้ Applicability Gate เพื่อตรวจสอบความเกี่ยวข้องกับภัยคุกคามทางไซเบอร์", first_indent=0.75)
    add_p(doc, "3. หากผลลัพธ์เป็น RETRIEVE ระบบจะเรียก Downstream RAG Service เพื่อดึงข้อมูลเทคนิค MITRE ATT&CK และบันทึกลง rag_contexts แต่หากเป็น SKIP ระบบจะข้ามไปโดยไม่เรียก RAG", first_indent=0.75)
    add_p(doc, "4. ระบบส่งคำสั่งไปยัง Main Case Analysis Engine เพื่อสร้างโครงสร้าง 5W1H ไทม์ไลน์ และรายการข้ออ้างเชิงประจักษ์", first_indent=0.75)
    add_p(doc, "5. Validation Engine ทำการตรวจสอบความถูกต้องของข้อความอ้างอิงตรงและพิกัดหน้าเอกสาร", first_indent=0.75)
    add_p(doc, "6. Follow-up Service ทำการวิเคราะห์ช่องว่างคดี (Gap Analysis) และตัดสินใจว่าจะถามคำถามหรือดำเนินการต่อ (PROCEED)", first_indent=0.75)
    add_p(doc, "7. Worker ทำการอัปเดตสถานะของ Run เป็น COMPLETED และปลดล็อกสัญญาเช่า", first_indent=0.75)

    add_h3(doc, "3.2.3 การคัดกรองความเกี่ยวข้องทางไซเบอร์ (Applicability Gate Flowchart)")
    add_p(doc, "ภาพที่ 3-4 แสดงขั้นตอนการคัดกรองความเกี่ยวข้องทางไซเบอร์ก่อนการสืบค้น (Pre-Retrieval Gate):")
    add_p(doc, "1. ตัวคัดกรองรับข้อมูลพยานหลักฐานจริงทั้งหมดในคดี", first_indent=0.75)
    add_p(doc, "2. โมเดลจำแนกพิจารณาว่าพยานหลักฐานมีพฤติการณ์โจมตีระบบคอมพิวเตอร์จริงหรือไม่ โดยใช้เกณฑ์ Precision-first", first_indent=0.75)
    add_p(doc, "3. หากพบเพียงพยานหลักฐานดิจิทัลทั่วไป เช่น การแชทผ่าน LINE สลิปโอนเงิน หรือกล้องวงจรปิด ประตูจะตัดสินเป็น 'SKIP' ทันที", first_indent=0.75)
    add_p(doc, "4. หากมีพฤติการณ์เจาะระบบ มัลแวร์ หรือการหลอกลวงขโมยเซสชัน ประตูจะตัดสินเป็น 'RETRIEVE' พร้อมระบุข้อความกระตุ้น (Trigger Span) ที่ตรงกับเอกสาร", first_indent=0.75)

    add_h3(doc, "3.2.4 การถามเพื่อความกระจ่างและตัดลูป (Clarification Flowchart)")
    add_p(doc, "ภาพที่ 3-5 แสดงขั้นตอนการจัดการคำถามขอความกระจ่างแบบมีสถานะ:")
    add_p(doc, "1. ระบบสกัดรายการช่องว่างพยานหลักฐานทั้งหมด (Gaps)", first_indent=0.75)
    add_p(doc, "2. กรองช่องว่างที่ไม่สามารถถามได้ (Unaskable) ออกไป", first_indent=0.75)
    add_p(doc, "3. กรองหัวข้อที่เคยถามไปแล้วและถูกบันทึกใน Exhausted Topics ออกไป", first_indent=0.75)
    add_p(doc, "4. หากไม่มีช่องว่างเหลือ ระบบจะจบกระบวนการและเลือก 'PROCEED' ทันที", first_indent=0.75)
    add_p(doc, "5. หากมีช่องว่างที่ถามได้ ระบบจะเลือกประเด็นที่มีความสำคัญสูงสุด (HIGH) และผูกกับข้อกล่าวหาหลัก มาสร้างคำถามที่กระชับ 1 คำถาม", first_indent=0.75)
    add_p(doc, "6. เมื่อผู้ใช้ส่งคำตอบกลับมา ระบบจะบันทึกหัวข้อนั้นลงในรายการ Exhausted Topics ทันที เพื่อไม่ให้ระบบถามซ้ำในประเด็นนี้อีกในรอบต่อไป", first_indent=0.75)

    add_h3(doc, "3.2.5 การสร้างรายงานสรุปสำนวนคดี (Report Generation Flowchart)")
    add_p(doc, "ภาพที่ 3-6 แสดงขั้นตอนการสร้างรายงานสรุปสำนวนคดีแบบกำหนดโครงสร้างแน่นอน:")
    add_p(doc, "1. ผู้ใช้กดปุ่ม 'Generate Report' จากหน้าจอรายงาน", first_indent=0.75)
    add_p(doc, "2. ระบบดึงข้อมูล Snapshot ล่าสุดของคดีที่ผ่านการตรวจสอบแล้วจากฐานข้อมูล", first_indent=0.75)
    add_p(doc, "3. ระบบประกอบข้อมูลเข้าสู่โครงสร้าง View Model ตามแม่แบบ 7 ส่วนมาตรฐาน", first_indent=0.75)
    add_p(doc, "4. เรียกใช้งานเอ็นจิน ReportLab เพื่อเรนเดอร์ข้อมูลเป็นไฟล์ PDF โดยตรงโดยไม่มีการเรียกโมเดลภาษาให้เขียนข้อความใหม่", first_indent=0.75)
    add_p(doc, "5. จัดเก็บไฟล์ PDF และบันทึกประวัติรายงานลงตาราง chat_reports พร้อมให้ผู้ใช้ดาวน์โหลด", first_indent=0.75)

    add_h2(doc, "3.3 Use Case Diagram ฟังก์ชันการทำงานของระบบ")
    add_p(doc, "การทำงานของระบบ CyberCase สามารถจำแนกบทบาทของผู้กระทำ (Actors) ออกเป็น 3 ส่วน ได้แก่:")
    add_p(doc, "1. พนักงานอัยการ / นิติกร (User/Prosecutor): ผู้ใช้งานหลักที่ทำหน้าที่นำเข้าสำนวนคดี ตรวจสอบภาพรวมคดี ตรวจสอบพยานหลักฐาน ตอบข้อสงสัย และสั่งพิมพ์รายงาน", first_indent=0.75)
    add_p(doc, "2. ระบบส่วนหลัง (Backend System): ผู้กระทำฝั่งระบบที่คอยจัดการการประมวลผล สกัดข้อเท็จจริง ตรวจสอบการอ้างอิง และจัดการการกู้คืนงาน", first_indent=0.75)
    add_p(doc, "3. บริการสืบค้นความรู้ภายนอก (Downstream RAG Service): บริการปลายน้ำที่ให้ข้อมูลทางเทคนิค MITRE ATT&CK", first_indent=0.75)
    
    p_fig_uc = doc.add_paragraph()
    p_fig_uc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig_uc.paragraph_format.space_before = Pt(8)
    p_fig_uc.paragraph_format.space_after = Pt(4)
    r_fuc = p_fig_uc.add_run("ภาพที่ 3-7 Use Case Diagram ฟังก์ชันการทำงานของระบบ CyberCase")
    r_fuc.bold = True

    add_h2(doc, "3.4 Use Case Description")
    add_p(doc, "เพื่อให้เห็นรายละเอียดขั้นตอนการทำงานอย่างชัดเจน จึงได้จัดทำตาราง Use Case Description สำหรับทุกฟังก์ชันการทำงานของระบบ CyberCase จำนวน 16 Use Case ดังนี้:")

    # Use Case 1
    add_use_case_table(
        doc, table_num="3-1", uc_num="1", uc_name="Create New Case Workspace",
        brief_desc="พนักงานอัยการทำการสร้างพื้นที่ทำงานสำหรับตรวจสำนวนคดีใหม่",
        actors="พนักงานอัยการ",
        pre_cond="ผู้ใช้เข้าสู่หน้าเว็บแอปพลิเคชัน CyberCase",
        post_cond="ระบบสร้างห้องสำนวนคดีใหม่และแสดงหน้านำเข้าเอกสาร",
        main_flow_actor="1. ผู้ใช้กดปุ่ม 'New Case' บนแถบเมนู",
        main_flow_system="1. ระบบส่งคำขอ POST /api/v1/chats ไปยังเซิร์ฟเวอร์\n2. เซิร์ฟเวอร์สร้างระเบียนใน chat_threads กำหนดสถานะเป็น 'idle'\n3. ระบบส่วนหน้านำผู้ใช้เข้าสู่หน้า Case Preparation Workspace",
        alt_flow="หากเซิร์ฟเวอร์ไม่สามารถเชื่อมต่อฐานข้อมูลได้ ระบบจะแสดงหน้าต่าง Meaningful Error Modal แจ้งเตือนข้อผิดพลาด",
        explanation="เมื่อพนักงานอัยการต้องการเริ่มตรวจพิจารณาสำนวนคดีใหม่ จะต้องกดสร้าง Workspace ซึ่งระบบจะกำหนดรหัสคดีเฉพาะ (Thread UUID) เพื่อใช้ในการอ้างอิงข้อมูลทั้งหมดของคดีนั้น"
    )

    # Use Case 2
    add_use_case_table(
        doc, table_num="3-2", uc_num="2", uc_name="Upload & Extract Case Document",
        brief_desc="พนักงานอัยการทำการอัปโหลดไฟล์เอกสารสำนวนคดีเพื่อสกัดข้อความ",
        actors="พนักงานอัยการ, ระบบส่วนหลัง",
        pre_cond="อยู่ในหน้า Case Preparation Workspace",
        post_cond="ระบบสกัดข้อความจากเอกสารและเตรียมเปิดหน้าต่างตรวจสอบ",
        main_flow_actor="1. ผู้ใช้เลือกไฟล์ PDF หรือ DOCX และกดปุ่ม Upload",
        main_flow_system="1. ระบบส่งไฟล์ไปยัง /api/v1/document-ingestion/preview\n2. ระบบทำการแยกหน้า สกัดข้อความ และคำนวณค่า OCR Confidence\n3. ส่งคืนโครงสร้างข้อความพร้อมเลขหน้ากลับมายังส่วนหน้า",
        alt_flow="หากไฟล์มีขนาดเกิน 25MB หรือเป็นรูปแบบไฟล์ที่ไม่รองรับ ระบบจะแจ้งเตือนและไม่อนุญาตให้อัปโหลด",
        explanation="ฟังก์ชันนี้เป็นการสกัดข้อความเบื้องต้นเพื่ออำนวยความสะดวกแก่ผู้ใช้ โดยข้อความที่สกัดได้ยังไม่ถือเป็นพยานหลักฐานจนกว่าผู้ใช้จะตรวจทานและกดยืนยัน"
    )

    # Use Case 3
    add_use_case_table(
        doc, table_num="3-3", uc_num="3", uc_name="Preview & Confirm Narrative",
        brief_desc="ผู้ใช้ตรวจทานข้อความที่สกัดได้จากเอกสารและกดยืนยันเพื่อเริ่มวิเคราะห์",
        actors="พนักงานอัยการ",
        pre_cond="ระบบสกัดข้อความสำเร็จและแสดงผลใน Preview Modal",
        post_cond="ข้อความได้รับการบันทึกลงฐานข้อมูลเป็นหลักฐานจริง (Evidence m_1)",
        main_flow_actor="1. ผู้ใช้อ่านตรวจทานข้อความ แก้ไขคำผิด\n2. กดปุ่ม 'Import & Analyze Case'",
        main_flow_system="1. บันทึกข้อความเป็น ChatMessage ในฐานข้อมูล\n2. คำนวณแฮช SHA-256 ของข้อความ\n3. เปลี่ยนสถานะ Thread เป็น 'processing' และสร้าง ChatRun",
        alt_flow="หากผู้ใช้กดยกเลิก ข้อความที่สกัดไว้จะถูกลบทิ้งโดยไม่มีการบันทึก",
        explanation="เป็นขั้นตอน Human-in-the-Loop เพื่อรับประกันว่าข้อความที่จะนำไปใช้วิเคราะห์ได้รับการตรวจสอบความถูกต้องจากมนุษย์แล้ว"
    )

    # Use Case 4
    add_use_case_table(
        doc, table_num="3-4", uc_num="4", uc_name="Execute Main Case Analysis",
        brief_desc="ระบบส่วนหลังดำเนินการประมวลผลวิเคราะห์สำนวนคดีในเบื้องหลัง",
        actors="ระบบส่วนหลัง",
        pre_cond="มี ChatRun ใหม่ในสถานะ 'queued'",
        post_cond="ระบบวิเคราะห์ข้อเท็จจริง สกัด 5W1H และบันทึก Trace สำเร็จ",
        main_flow_actor="ผู้ใช้รอรับผลการประมวลผลผ่านหน้าจอแสดงสถานะ",
        main_flow_system="1. Worker เคลมสัญญาเช่า Run และส่ง Heartbeat\n2. ส่ง Prompt พร้อมข้อความพยานหลักฐานไปยัง LLM\n3. รับคำตอบและถอดรหัสโครงสร้าง Pydantic AnalysisTraceV3\n4. บันทึกผลลัพธ์และเปลี่ยนสถานะเป็น 'awaiting_followup' หรือ 'answered'",
        alt_flow="หากการประมวลผลเกิน 6 นาทีโดยไม่มี Heartbeat ระบบ Recovery จะทำเครื่องหมายว่า 'interrupted'",
        explanation="การทำงานหลักของระบบในการสกัดข้อเท็จจริงทางคดี ซึ่งประมวลผลแบบ Asynchronous ในเบื้องหลังเพื่อไม่ให้บล็อกการทำงานของหน้าเว็บ"
    )

    # Use Case 5
    add_use_case_table(
        doc, table_num="3-5", uc_num="5", uc_name="Classify Technical Applicability",
        brief_desc="ระบบคัดกรองว่าพยานหลักฐานในคดีเกี่ยวข้องกับภัยคุกคามทางไซเบอร์หรือไม่",
        actors="ระบบส่วนหลัง",
        pre_cond="อยู่ในขั้นตอนเริ่มประมวลผล Run",
        post_cond="ได้ผลการตัดสินใจเป็น RETRIEVE หรือ SKIP",
        main_flow_actor="-",
        main_flow_system="1. ส่งข้อความหลักฐานไปยังโมเดลจำแนก Applicability Gate\n2. ตรวจสอบพฤติการณ์โจมตีทางเทคนิค\n3. หากเป็นพยานหลักฐานทั่วไป (LINE, โอนเงิน) ให้ตัดสินเป็น SKIP\n4. หากมีการโจมตีระบบ ให้ระบุ Trigger Span และตัดสินเป็น RETRIEVE",
        alt_flow="หากเกิดข้อผิดพลาดในการประมวลผล ให้ใช้ค่าเริ่มต้นเป็น SKIP เพื่อความปลอดภัย",
        explanation="ป้องกันไม่ให้ระบบนำเอาเทคนิค MITRE ATT&CK มาจับคู่กับคดีอาญาทั่วไปโดยไม่จำเป็น"
    )

    # Use Case 6
    add_use_case_table(
        doc, table_num="3-6", uc_num="6", uc_name="Request Downstream Technical Intelligence",
        brief_desc="ระบบเรียกข้อมูลเทคนิคจากบริการ GraphRAG ภายนอกเมื่อคดีเข้าข่ายไซเบอร์",
        actors="ระบบส่วนหลัง, Downstream RAG Service",
        pre_cond="ผลการคัดกรอง Applicability Gate เป็น 'RETRIEVE'",
        post_cond="ข้อมูล MITRE ATT&CK ถูกบันทึกลงตาราง rag_contexts",
        main_flow_actor="-",
        main_flow_system="1. ส่งคำขอ HTTP POST ไปยังบริการ GraphRAG พร้อมข้อความ Trigger\n2. รับข้อมูลเทคนิค ยุทธวิธี และแนวทางแก้ไขกลับมา\n3. ตรวจสอบความถูกต้องของโครงสร้าง JSON และบันทึกลงฐานข้อมูล",
        alt_flow="หากบริการภายนอกไม่ตอบสนอง ระบบจะบันทึกสถานะว่า Context Unavailable และดำเนินการวิเคราะห์คดีต่อไปได้โดยไม่หยุดชะงัก",
        explanation="เป็นการดึงองค์ความรู้ทางเทคนิคมาเสริม โดยกักกันข้อมูลไว้ในชั้นบริบทและไม่ให้ปะปนกับพยานหลักฐานของคดี"
    )

    # Use Case 7
    add_use_case_table(
        doc, table_num="3-7", uc_num="7", uc_name="Extract 5W1H Overview & Chronology",
        brief_desc="ระบบสกัดโครงสร้าง 5W1H และเรียงลำดับเหตุการณ์ตามวันเวลาจริง",
        actors="ระบบส่วนหลัง",
        pre_cond="โมเดลภาษาประมวลผลข้อความสำนวนคดี",
        post_cond="ได้บทสรุปภาพรวม ผู้เกี่ยวข้อง และไทม์ไลน์เหตุการณ์",
        main_flow_actor="-",
        main_flow_system="1. สกัดตัวตนผู้เสียหาย ผู้ต้องหา พยาน และเจ้าหน้าที่\n2. สกัดวันเวลาและข้อความเหตุการณ์\n3. จัดเรียงลำดับเหตุการณ์ตามเวลาที่ระบุจริง มิใช่ลำดับที่ปรากฏในเอกสาร",
        alt_flow="หากในสำนวนไม่ระบุวันเวลาที่แน่ชัด ระบบจะจัดกลุ่มเป็นเหตุการณ์ที่ไม่ระบุเวลา",
        explanation="ช่วยให้พนักงานอัยการสามารถมองเห็นเส้นเวลาของการกระทำความผิดได้อย่างเป็นลำดับและถูกต้องตามความเป็นจริง"
    )

    # Use Case 8
    add_use_case_table(
        doc, table_num="3-8", uc_num="8", uc_name="Categorize Epistemic Claims",
        brief_desc="ระบบจำแนกข้ออ้างแต่ละข้อออกตามสถานะเชิงประจักษ์และความไม่แน่นอน",
        actors="ระบบส่วนหลัง",
        pre_cond="สกัดข้อความข้อกล่าวหาและพฤติการณ์คดี",
        post_cond="ข้ออ้างแต่ละข้อถูกกำหนดสถานะ (Reported, Inference, Unresolved, Disputed)",
        main_flow_actor="-",
        main_flow_system="1. กำหนดรหัสข้ออ้างแบบคงที่ (A-01 ถึง A-64)\n2. แยกข้อความที่เป็นคำให้การพยาน (Reported) ออกจากข้ออนุมานของพนักงานสอบสวน (Inference)\n3. หากมีคำให้การขัดแย้งกันให้กำหนดเป็น Disputed",
        alt_flow="หากโมเดลไม่ระบุสถานะ ระบบจะกำหนดค่าเริ่มต้นเป็น Unresolved เพื่อความปลอดภัย",
        explanation="หัวใจสำคัญในการป้องกันไม่ให้พนักงานอัยการสับสนระหว่างข้อเท็จจริงที่มีหลักฐานยืนยันกับข้ออนุมานหรือความเห็นลอย ๆ"
    )

    # Use Case 9
    add_use_case_table(
        doc, table_num="3-9", uc_num="9", uc_name="Validate Verbatim Source Quotes",
        brief_desc="ระบบตรวจสอบว่าข้อความอ้างอิงตรงจากสำนวนมีอยู่จริงในเอกสารต้นฉบับ",
        actors="ระบบส่วนหลัง",
        pre_cond="โมเดลส่งข้ออ้างพร้อมข้อความอ้างอิง (Quotation Candidate)",
        post_cond="ข้อความอ้างอิงที่ถูกต้องได้รับการอนุมัติและผูกกับพิกัดหน้า",
        main_flow_actor="-",
        main_flow_system="1. ค้นหาข้อความอ้างอิงในเนื้อหาหลักฐานต้นทาง (m_1)\n2. ตรวจสอบว่าตรงกันแบบตัวอักษรต่อตัวอักษร (Verbatim Substring)\n3. หากตรงและระบุหน้าชัดเจน ให้บันทึกพิกัดหน้า\n4. หากตรงแต่หน้าคลุมเครือ ให้ลดระดับเหลือเพียงการอ้างอิงเอกสาร\n5. หากข้อความไม่ตรง ให้ตัดการอ้างอิงทิ้งทันที",
        alt_flow="หากโมเดลสร้างข้อความอ้างอิงปลอม ระบบจะตัดสิทธิ์การอ้างอิงและบันทึกข้อผิดพลาดใน Trace",
        explanation="ขจัดปัญหาการแต่งข้อความอ้างอิงของโมเดลภาษา ทำให้มั่นใจได้ว่าข้อความที่ไฮไลต์ในเอกสารเป็นของจริง 100%"
    )

    # Use Case 10
    add_use_case_table(
        doc, table_num="3-10", uc_num="10", uc_name="Detect Evidentiary Gaps",
        brief_desc="ระบบวิเคราะห์หาจุดบกพร่องหรือช่องว่างของพยานหลักฐานที่ยังขาดหายไปในสำนวน",
        actors="ระบบส่วนหลัง",
        pre_cond="สร้างข้ออ้างทางคดีเรียบร้อยแล้ว",
        post_cond="ได้รายการประเด็นช่องว่างคดี (Gaps) พร้อมระดับความสำคัญ (HIGH/MEDIUM/LOW)",
        main_flow_actor="-",
        main_flow_system="1. ตรวจสอบความสมบูรณ์ขององค์ประกอบความผิดและพยานหลักฐานสนับสนุน\n2. ระบุประเด็นที่ขาด เช่น ขาดผลชันสูตร, ขาดการตรวจลายนิ้วมือ, ขาดสลิปยืนยันยอด\n3. ผูกช่องว่างเข้ากับข้ออ้างที่ได้รับผลกระทบ (affected_claim_ids)",
        alt_flow="หากสำนวนมีพยานหลักฐานครบถ้วนสมบูรณ์ รายการช่องว่างจะเป็นค่าว่าง",
        explanation="ช่วยพนักงานอัยการในการตรวจหาจุดอ่อนของสำนวนเพื่อเตรียมสั่งสอบสวนเพิ่มเติมได้อย่างตรงจุด"
    )

    # Use Case 11
    add_use_case_table(
        doc, table_num="3-11", uc_num="11", uc_name="Select & Formulate Clarification Question",
        brief_desc="ระบบคัดเลือกประเด็นช่องว่างที่สำคัญที่สุดมาสร้างคำถามขอความกระจ่าง 1 ข้อ",
        actors="ระบบส่วนหลัง",
        pre_cond="มีรายการช่องว่างคดีที่สามารถถามได้ (Askable Gaps)",
        post_cond="ได้คำถามที่ตรงเป้าหมาย 1 คำถาม หรือตัดสินใจดำเนินการต่อ (PROCEED)",
        main_flow_actor="-",
        main_flow_system="1. กรองประเด็นที่เคยถามไปแล้วในอดีตออก\n2. เรียงลำดับความสำคัญ (HIGH มาก่อน MEDIUM)\n3. เลือกประเด็นที่มีการผูกกับข้อกล่าวหาหลัก\n4. สั่งให้โมเดลสร้างคำถามที่กระชับ ชัดเจน และตรงประเด็น 1 คำถาม",
        alt_flow="หากไม่มีประเด็นที่ถามได้เหลืออยู่ ระบบจะออกคำสั่ง PROCEED และเข้าสู่สถานะวิเคราะห์เสร็จสมบูรณ์",
        explanation="จำกัดการถามให้มีประสิทธิภาพสูงสุด ไม่ถามสะเปะสะปะ และถามเฉพาะเรื่องสำคัญต่อรูปคดี"
    )

    # Use Case 12
    add_use_case_table(
        doc, table_num="3-12", uc_num="12", uc_name="Exhaust Clarification Topic",
        brief_desc="ระบบบันทึกปิดประเด็นคำถามเมื่อผู้ใช้ตอบแล้ว เพื่อป้องกันการถามวนซ้ำ",
        actors="ระบบส่วนหลัง",
        pre_cond="ผู้ใช้ส่งคำตอบสำหรับคำถามขอความกระจ่างเข้ามา",
        post_cond="หัวข้อคำถาม (Normalized Topic) ถูกเพิ่มลงในรายการ Exhausted Topics",
        main_flow_actor="ผู้ใช้พิมพ์คำชี้แจงหรือข้อมูลเพิ่มเติมและกดส่ง",
        main_flow_system="1. นำหัวข้อคำถามมาแปลงเป็นรูปมาตรฐาน (Lowercase, Underscore)\n2. เพิ่มหัวข้อลงในประวัติคำถามที่สิ้นสุดแล้ว\n3. บันทึกคำตอบของผู้ใช้เป็นพยานหลักฐานใหม่ (m_2)\n4. สั่งเริ่มการวิเคราะห์รอบใหม่เพื่อปรับปรุงภาพรวมคดี",
        alt_flow="-",
        explanation="กลไกสำคัญในการตัดลูปการสนทนา ทำให้ระบบไม่ถามซ้ำในเรื่องเดิม แม้ว่าข้อมูลจะยังไม่ครบสมบูรณ์แต่ระบบจะบันทึกเป็นข้อจำกัดแทน"
    )

    # Use Case 13
    add_use_case_table(
        doc, table_num="3-13", uc_num="13", uc_name="Inspect Source Passage in Drawer",
        brief_desc="พนักงานอัยการตรวจสอบข้อความจริงในสำนวนคดีผ่านลิ้นชักเอกสาร",
        actors="พนักงานอัยการ",
        pre_cond="อยู่ในหน้า Case Overview และมีชิปอ้างอิงแหล่งที่มา",
        post_cond="ลิ้นชักเปิดออกและเลื่อนไปยังข้อความที่ตรงกันพร้อมไฮไลต์สีเหลือง",
        main_flow_actor="1. ผู้ใช้คลิกที่ชิปอ้างอิงแหล่งที่มา (เช่น [Page 3]) บนการ์ดข้ออ้าง",
        main_flow_system="1. ระบบส่วนหน้าเปิดลิ้นชัก Citation Drawer\n2. แสดงเนื้อหาเอกสารสำนวนคดีในหน้าที่เกี่ยวข้อง\n3. เลื่อนหน้าจอ (Scroll) ไปยังตำแหน่งข้อความและทำไฮไลต์ข้อความอ้างอิงตรงทันที",
        alt_flow="หากเป็นการอ้างอิงระดับเอกสารที่ไม่ระบุหน้า ระบบจะแสดงต้นฉบับและไฮไลต์ข้อความในจุดที่พบข้อความตรงกัน",
        explanation="ช่วยให้อัยการสามารถพิสูจน์ข้อเท็จจริงได้ด้วยตาตนเองภายในเวลาไม่กี่วินาทีโดยไม่ต้องเปิดหาในแฟ้มกระดาษ"
    )

    # Use Case 14
    add_use_case_table(
        doc, table_num="3-14", uc_num="14", uc_name="Ask Contextual Question",
        brief_desc="พนักงานอัยการสอบถามข้อสงสัยเพิ่มเติมต่อสำนวนคดีในโหมดซักถาม",
        actors="พนักงานอัยการ",
        pre_cond="การวิเคราะห์สำนวนคดีหลักเสร็จสมบูรณ์แล้ว",
        post_cond="ระบบตอบคำถามโดยอ้างอิงจากบทวิเคราะห์เดิมโดยไม่สร้างหลักฐานใหม่",
        main_flow_actor="1. ผู้ใช้พิมพ์คำถามลงในช่องสนทนา เช่น 'พยานปากที่ 2 ยืนยันเรื่องเวลาอย่างไร'",
        main_flow_system="1. ระบบเรียก API ในโหมด 'ask'\n2. ดึง Snapshot บทวิเคราะห์และบริบทเดิมมาตอบคำถาม\n3. ส่งคำตอบให้ผู้ใช้โดยไม่บันทึกคำถามเป็นพยานหลักฐานและไม่เรียก RAG ใหม่",
        alt_flow="หากคำถามอยู่นอกเหนือข้อเท็จจริงในสำนวน ระบบจะแจ้งว่าไม่มีข้อมูลปรากฏในเอกสาร",
        explanation="แยกแยะอย่างชัดเจนระหว่าง 'การเพิ่มพยานหลักฐานคดี' กับ 'การซักถามทำความเข้าใจของผู้ใช้'"
    )

    # Use Case 15
    add_use_case_table(
        doc, table_num="3-15", uc_num="15", uc_name="Generate Deterministic Forensic Report",
        brief_desc="ระบบสร้างรายงานสรุปสำนวนการสอบสวนอย่างเป็นทางการในรูปแบบ PDF",
        actors="พนักงานอัยการ, ระบบส่วนหลัง",
        pre_cond="การวิเคราะห์คดีเสร็จสมบูรณ์และผู้ใช้กดปุ่มสร้างรายงาน",
        post_cond="ได้ไฟล์เอกสาร PDF รายงานสรุปสำนวนคดีมาตรฐาน 7 ส่วน",
        main_flow_actor="1. ผู้ใช้กดปุ่ม 'Generate Report' หรือ 'Download PDF'",
        main_flow_system="1. ดึง Snapshot ผลการวิเคราะห์ล่าสุดจากฐานข้อมูล\n2. ประกอบข้อมูลเข้าสู่แม่แบบ 7 ส่วนผ่าน View Model Builder\n3. สั่ง ReportLab เรนเดอร์เป็นไฟล์ PDF ตามพิกัดและฟอนต์มาตรฐาน\n4. ส่งไฟล์ PDF ให้ผู้ใช้ดาวน์โหลด",
        alt_flow="หาก Snapshot เสียหายหรือไม่สมบูรณ์ ระบบจะปฏิเสธการสร้างรายงานและแจ้งเตือนให้ทำการวิเคราะห์ใหม่",
        explanation="รับประกันว่าเอกสารรายงานที่ออกโดยระบบจะตรงกับสิ่งที่อัยการเห็นบนหน้าจอทุกประการ โดยไม่มีการสร้างข้อความแปลกปลอมขึ้นมาใหม่"
    )

    # Use Case 16
    add_use_case_table(
        doc, table_num="3-16", uc_num="16", uc_name="Retry Interrupted Processing Run",
        brief_desc="ผู้ใช้สั่งกู้คืนและประมวลผลซ้ำสำหรับงานที่หยุดชะงักจากปัญหาเครือข่าย",
        actors="พนักงานอัยการ, ระบบส่วนหลัง",
        pre_cond="Run มีสถานะเป็น 'interrupted' หรือ 'failed'",
        post_cond="ระบบเริ่มการประมวลผลใหม่โดยใช้ข้อมูลและหลักฐานเดิม",
        main_flow_actor="1. ผู้ใช้กดปุ่ม 'Retry Analysis' บนการ์ดแจ้งเตือน",
        main_flow_system="1. ระบบตรวจสอบลายนิ้วมือคำขอเดิม (Request Fingerprint)\n2. สร้าง Run ใหม่ภายใต้ Thread เดิมโดยใช้ Evidence Message เดิม\n3. ป้องกันการสร้างข้อความหลักฐานซ้ำซ้อน\n4. เริ่มต้นการทำงานของ Worker เพื่อวิเคราะห์ต่อ",
        alt_flow="หากมีข้อความใหม่เกิดขึ้นหลังจาก Run ที่ขัดข้อง ระบบจะไม่อนุญาตให้กู้คืน Run เก่าเพื่อป้องกันความขัดแย้งของข้อมูล",
        explanation="เพิ่มความทนทานของระบบเมื่อเกิดปัญหาเซิร์ฟเวอร์ขัดข้องหรือเครือข่ายหลุดระหว่างที่โมเดลกำลังประมวลผล"
    )

    add_h2(doc, "3.5 Entity-Relationship Diagram (ER Diagram)")
    add_p(doc, "โครงสร้างฐานข้อมูลของระบบ CyberCase ได้รับการออกแบบบนระบบจัดการฐานข้อมูลเชิงสัมพันธ์ PostgreSQL โดยใช้มาตรฐานรหัสเฉพาะสากล (UUID v4) ในการระบุเอกลักษณ์ของแต่ละระเบียน และใช้ข้อมูลประเภท JSONB ในการจัดเก็บข้อมูลที่มีโครงสร้างยืดหยุ่น เช่น Metadata และ Analysis Trace ดังแสดงความสัมพันธ์ในภาพที่ 3-8")
    
    p_fig_er = doc.add_paragraph()
    p_fig_er.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig_er.paragraph_format.space_before = Pt(8)
    p_fig_er.paragraph_format.space_after = Pt(4)
    r_fer = p_fig_er.add_run("ภาพที่ 3-8 Entity-Relationship Diagram (ER Diagram) โครงสร้างฐานข้อมูล CyberCase")
    r_fer.bold = True

    add_h2(doc, "3.6 Data Dictionary")
    add_p(doc, "พจนานุกรมข้อมูล (Data Dictionary) แสดงรายละเอียดของโครงสร้างตาราง ชนิดข้อมูล ข้อจำกัด และคำอธิบายความหมายของแต่ละคอลัมน์ในฐานข้อมูล PostgreSQL จำนวน 5 ตารางหลัก ดังนี้:")

    # Table 3-17 chat_threads
    add_data_dict_table(
        doc, table_num="3-17", table_name="chat_threads",
        table_desc="ข้อมูลห้องสำนวนคดีและพื้นที่ทำงานสำหรับตรวจพิจารณาสำนวน",
        rows_data=[
            ["PK", "id", "UUID", "NOT NULL", "รหัสประจำห้องสำนวนคดี (UUID v4)", "-"],
            ["-", "title", "varchar(255)", "NOT NULL", "ชื่อสำนวนคดีหรือหัวข้อคดี", "-"],
            ["-", "status", "varchar(24)", "NOT NULL", "สถานะห้อง (idle, processing, awaiting_followup, answered, failed)", "-"],
            ["-", "next_message_ordinal", "integer", "NOT NULL", "ลำดับข้อความถัดไปในห้อง (เริ่มต้นที่ 1)", "-"],
            ["-", "created_at", "timestamptz", "NOT NULL", "วันเวลาที่สร้างห้องสำนวนคดี", "-"],
            ["-", "updated_at", "timestamptz", "NOT NULL", "วันเวลาที่มีการแก้ไขข้อมูลล่าสุด", "-"]
        ]
    )

    # Table 3-18 chat_messages
    add_data_dict_table(
        doc, table_num="3-18", table_name="chat_messages",
        table_desc="ข้อความพยานหลักฐาน ข้อความชี้แจง และผลการวิเคราะห์ของระบบ",
        rows_data=[
            ["PK", "id", "UUID", "NOT NULL", "รหัสประจำข้อความ (UUID v4)", "-"],
            ["FK", "thread_id", "UUID", "NOT NULL", "รหัสห้องสำนวนคดีที่สังกัด", "chat_threads.id"],
            ["-", "role", "varchar(24)", "NOT NULL", "บทบาทผู้ส่ง (user, assistant, system)", "-"],
            ["-", "content", "text", "NOT NULL", "เนื้อหาข้อความหรือบทวิเคราะห์", "-"],
            ["-", "ordinal", "integer", "NOT NULL", "ลำดับข้อความภายในห้องสำนวน", "-"],
            ["-", "status", "varchar(24)", "NOT NULL", "สถานะข้อความ (active, edited, deleted)", "-"],
            ["-", "evidence_hash", "varchar(64)", "NULL", "ค่าแฮช SHA-256 ของพยานหลักฐาน", "-"],
            ["-", "metadata_", "jsonb", "NOT NULL", "ข้อมูลบริบท, ผลลัพธ์ Trace v3, การอ้างอิง", "-"],
            ["-", "created_at", "timestamptz", "NOT NULL", "วันเวลาที่บันทึกข้อความ", "-"],
            ["-", "updated_at", "timestamptz", "NOT NULL", "วันเวลาที่มีการแก้ไขล่าสุด", "-"]
        ]
    )

    # Table 3-19 chat_runs
    add_data_dict_table(
        doc, table_num="3-19", table_name="chat_runs",
        table_desc="รอบการประมวลผลเบื้องหลังและการถือครองสัญญาเช่าของ Worker",
        rows_data=[
            ["PK", "id", "UUID", "NOT NULL", "รหัสประจำรอบการประมวลผล (UUID v4)", "-"],
            ["FK", "thread_id", "UUID", "NOT NULL", "รหัสห้องสำนวนคดีที่ประมวลผล", "chat_threads.id"],
            ["FK", "trigger_message_id", "UUID", "NOT NULL", "รหัสข้อความหลักฐานที่กระตุ้นให้รัน", "chat_messages.id"],
            ["-", "action", "varchar(32)", "NOT NULL", "การกระทำ (initial, clarification, ask, add_info)", "-"],
            ["-", "status", "varchar(24)", "NOT NULL", "สถานะรอบการรัน (queued, running, completed, failed, interrupted)", "-"],
            ["-", "lease_owner", "varchar(128)", "NULL", "รหัสระบุตัวตนของ Worker ที่ถือครองงาน", "-"],
            ["-", "lease_expires_at", "timestamptz", "NULL", "วันเวลาหมดอายุของสัญญาเช่า (ต่ออายุทุก 30 วินาที)", "-"],
            ["-", "error_details", "jsonb", "NULL", "รายละเอียดข้อผิดพลาดหากการรันล้มเหลว", "-"],
            ["-", "retry_count", "integer", "NOT NULL", "จำนวนครั้งที่มีการลองใหม่ (Default 0)", "-"],
            ["-", "created_at", "timestamptz", "NOT NULL", "วันเวลาที่เริ่มสร้างคำขอรัน", "-"],
            ["-", "updated_at", "timestamptz", "NOT NULL", "วันเวลาที่อัปเดตสถานะล่าสุด", "-"]
        ]
    )

    # Table 3-20 rag_contexts
    add_data_dict_table(
        doc, table_num="3-20", table_name="rag_contexts",
        table_desc="ข้อมูลบริบทองค์ความรู้ภายนอก MITRE ATT&CK ที่สืบค้นได้จาก RAG",
        rows_data=[
            ["PK", "id", "UUID", "NOT NULL", "รหัสประจำบริบท RAG (UUID v4)", "-"],
            ["FK", "thread_id", "UUID", "NOT NULL", "รหัสห้องสำนวนคดี", "chat_threads.id"],
            ["FK", "run_id", "UUID", "NOT NULL", "รหัสรอบการรันที่ทำการสืบค้น", "chat_runs.id"],
            ["-", "query_text", "text", "NOT NULL", "ข้อความคำค้นที่ส่งไปยังบริการ RAG", "-"],
            ["-", "retrieved_context", "text", "NOT NULL", "เนื้อหาองค์ความรู้ที่ได้รับตอบกลับมา", "-"],
            ["-", "techniques", "jsonb", "NOT NULL", "รายการเทคนิค MITRE ATT&CK ที่สอดคล้อง", "-"],
            ["-", "similarity_score", "float", "NULL", "คะแนนความคล้ายคลึงเชิงความหมาย", "-"],
            ["-", "created_at", "timestamptz", "NOT NULL", "วันเวลาที่บันทึกข้อมูลบริบท", "-"]
        ]
    )

    # Table 3-21 chat_reports
    add_data_dict_table(
        doc, table_num="3-21", table_name="chat_reports",
        table_desc="รายงานสรุปสำนวนการสอบสวนอย่างเป็นทางการในรูปแบบเอกสาร",
        rows_data=[
            ["PK", "id", "UUID", "NOT NULL", "รหัสประจำรายงาน (UUID v4)", "-"],
            ["FK", "thread_id", "UUID", "NOT NULL", "รหัสห้องสำนวนคดี", "chat_threads.id"],
            ["FK", "run_id", "UUID", "NOT NULL", "รหัสรอบการรันที่ใช้เป็นข้อมูลต้นร่าง", "chat_runs.id"],
            ["-", "version", "integer", "NOT NULL", "หมายเลขเวอร์ชันของรายงาน (1, 2, ...)", "-"],
            ["-", "title", "varchar(255)", "NOT NULL", "ชื่อเอกสารรายงานสรุปสำนวนคดี", "-"],
            ["-", "executive_summary", "text", "NOT NULL", "บทสรุปผู้บริหารและภาพรวม 5W1H", "-"],
            ["-", "markdown_content", "text", "NOT NULL", "เนื้อหารายงานฉบับเต็มในรูปแบบ Markdown", "-"],
            ["-", "pdf_blob", "bytea", "NULL", "ข้อมูลไบนารีของไฟล์ PDF ที่เรนเดอร์สำเร็จ", "-"],
            ["-", "created_at", "timestamptz", "NOT NULL", "วันเวลาที่สร้างรายงาน", "-"],
            ["-", "updated_at", "timestamptz", "NOT NULL", "วันเวลาที่มีการปรับปรุงล่าสุด", "-"]
        ]
    )

    doc.add_page_break()

print("Chapter 3 ready.")
